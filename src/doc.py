from docx import Document
from docx.shared import Inches, Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.style import WD_STYLE_TYPE

from docx.compat import Unicode
from docx.oxml import OxmlElement
from docx.oxml.shared import qn

from docx.oxml.exceptions import InvalidXmlError
from docx.oxml.ns import NamespacePrefixedTag, nsmap, qn
from docx.shared import lazyproperty

from lxml import etree


import lorem
from lorem.text import TextLorem

class Documento:

    teste = """A1;*B1*;C1
A2;B2;C2
A3 isso é um teste do tamanho da coluna com um texto grande ;B3;C3
AB4;;C4"""

    merge_teste = {
        "linha": 3,
        "colunas": [0, 1]
    }

    arquivo = ""

    def __init__(self, arquivo):
        self.arquivo = arquivo
        self.modelo = Document("resources/modelo.docx")
        allTables = self.modelo.tables
        for activeTable in allTables:
            activeTable._element.getparent().remove(activeTable._element)
        self.modelo.save('temp.docx')

        self.documento = Document('temp.docx')
        styles = self.documento.styles
        titulo = styles.add_style('Titulo', WD_STYLE_TYPE.PARAGRAPH)
        titulo.base_style = styles['Normal']
        convoc = styles.add_style('Convocacao', WD_STYLE_TYPE.PARAGRAPH)
        convoc.base_style = styles['Normal']
        pauta = styles.add_style('Pauta', WD_STYLE_TYPE.PARAGRAPH)
        pauta.base_style = styles['Normal']
        data = styles.add_style('Data', WD_STYLE_TYPE.PARAGRAPH)
        data.base_style = styles['Normal']
        assinatura = styles.add_style('Assinatura', WD_STYLE_TYPE.PARAGRAPH)
        assinatura.base_style = styles['Normal']
        rodape = styles.add_style('Rodape', WD_STYLE_TYPE.PARAGRAPH)
        rodape.base_style = styles['Normal']
        ata = styles.add_style('Ata', WD_STYLE_TYPE.PARAGRAPH)
        ata.base_style = styles['Normal']
        # pauta.base_style = styles['Normal']

        paragraph_styles = [s for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH]
        for style in paragraph_styles:
            print(style.name)

        # self.tabela(
        #     dados=self.teste,
        #     merge=[self.merge_teste],
        #     tamanho=[12, 4, 1]
        # )
        #
        # self.documento.add_paragraph("Teste")
        # self.add_conceito("Cocneito teste")
        # self.documento.add_paragraph()
        # self.add_caixa("Código")
        # self.documento.add_paragraph()
        # self.add_assinatura("NOME", "TESTE")


    def salvar(self):
        self.documento.save(self.arquivo)

    def cabecalho(self, departamento):

        section = self.documento.sections[0]
        header = section.header

        tabela = header.add_table(1, 3, Inches(6))

        celulas = tabela.rows[0].cells

        celulas[0].width = Cm(3)
        p = celulas[0].paragraphs[0]
        r = p.add_run()
        r.add_picture('resources/ufrgs.png', width=Cm(2.7))
        celulas[0].vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM

        departamento = "departamento de " + departamento
        departamento = departamento.upper()

        celulas[1].text = 'INSTITUTO DE GEOCIÊNCIAS\n' + departamento
        celulas[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        celulas[1].vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM

        celulas[2].width = Cm(3)
        p = celulas[2].paragraphs[0]
        r = p.add_run()
        r.add_picture('resources/igeo.png', width=Cm(2.7))
        celulas[2].vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM

    def numerar(self):
        section = self.documento.sections[-1]
        sectPr = section._sectPr
        lnNumType = OxmlElement('w:lnNumType')
        lnNumType.set(qn('w:countBy'), '1')
        lnNumType.set(qn('w:restart'), 'newSection')
        sectPr.insert_element_before(lnNumType, 'w:cols')

    def titulo(self, texto):
        self.paragrafo(texto, "Titulo",
                       fonte=16,
                       bold=True,
                       alinhamento="Centralizado"
                       )
        # p = self.documento.add_paragraph()
        # paragraph_format = self.documento.styles['Titulo'].paragraph_format
        # paragraph_format.first_line_indent = Pt(0)
        # p.style = self.documento.styles['Titulo']
        # p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # r = p.add_run(texto)
        # r.bold = True
        # r.font.size = Pt(16)

    def ata(self, abertura, pautas, encerramento):
        self.numerar()
        p = self.documento.add_paragraph()
        paragraph_format = self.documento.styles["Ata"].paragraph_format
        paragraph_format.left_indent = Pt(5)
        p.style = self.documento.styles["Ata"]
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(abertura)
        cont = 0
        for item in pautas:
            cont += 1
            r = p.add_run(str(cont) + ") " + item + ":")
            r.bold = True
            r = p.add_run(" PAUTA] ")
            r.bold = False
        p.add_run(encerramento)

    def convoc(self, texto):
        self.paragrafo(texto, "Convocacao", recuo=80)
        # p = self.documento.add_paragraph()
        # paragraph_format = self.documento.styles['Convocacao'].paragraph_format
        # paragraph_format.first_line_indent = Pt(42)
        # p.style = self.documento.styles['Titulo']
        # p.add_run(texto)

    def pauta(self, texto):
        p = self.documento.add_paragraph()
        paragraph_format = self.documento.styles["Pauta"].paragraph_format
        # paragraph_format.first_line_indent = Pt(80)
        paragraph_format.left_indent = Pt(42)
        p.style = self.documento.styles["Pauta"]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        cont = 0
        for item in texto:
            cont += 1
            r = p.add_run("\t\t" + str(cont) + ". " + item + "\n")

    def data(self, texto):
        self.paragrafo(texto, "Data", alinhamento="Direita")

    def assinatura(self, texto):
        self.paragrafo(texto, "Assinatura", alinhamento="Centralizado")

    def rodape(self, prof, tecnicos, alunos):
        self.paragrafo("\nPrezado(a) Membro do Departamento:", "Data", alinhamento="Esquerda")
        # p = self.documento.add_paragraph()
        # p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # p.add_run("Prezado(a) Membro do Departamento:\n")
        # p = self.documento.add_paragraph()
        # p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph_format = self.documento.styles["Rodape"].paragraph_format
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)
        for texto, lista in [
            ("Docentes", prof),
            ("Técnicos", tecnicos),
            ("Discentes", alunos)
        ]:
            if len(lista):
                p = self.documento.add_paragraph()
                p.style = self.documento.styles["Rodape"]
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                r = p.add_run(texto)
                r.italic = True
                r.underline = True
                p.add_run(": ")
                p.add_run(lista)
                p.add_run(".")

    def paragrafo(self,
                  texto,
                  estilo,
                  recuo=0,
                  fonte=12,
                  bold=False,
                  alinhamento="Justificado"
                  ):
        p = self.documento.add_paragraph()
        paragraph_format = self.documento.styles[estilo].paragraph_format
        paragraph_format.first_line_indent = Pt(recuo)
        p.style = self.documento.styles[estilo]
        if alinhamento == "Justificado":
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif alinhamento == "Centralizado":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alinhamento == "Direita":
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(texto)
        r.bold = bold
        r.font.size = Pt(fonte)

